# from flask import Flask, render_template, request, jsonify, session
# from werkzeug.utils import secure_filename
# import os
# from pathlib import Path
# import tempfile
# import logging

# from dotenv import load_dotenv
# import pandas as pd
# import docx
# import pypdf

# from groq import Groq
# from llama_index.core import VectorStoreIndex, Document
# from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# # Load environment variables
# load_dotenv(".env.local")
# GROQ_API_KEY = os.getenv("GROQ_API_KEY")
# if not GROQ_API_KEY:
#     raise ValueError("Missing GROQ_API_KEY in your .env.local file")

# # Configure logging
# logging.basicConfig(level=logging.INFO)

# # Flask app setup
# app = Flask(__name__)
# app.secret_key = os.urandom(24)
# app.config['UPLOAD_FOLDER'] = 'uploads'
# app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB

# Path(app.config['UPLOAD_FOLDER']).mkdir(parents=True, exist_ok=True)

# # Groq and Embedding model setup
# groq_client = Groq(api_key=GROQ_API_KEY)
# embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-mpnet-base-v2")

# # Utility: Process files to extract text
# def process_file(file_path):
#     ext = Path(file_path).suffix.lower()
#     try:
#         if ext == '.pdf':
#             with open(file_path, 'rb') as f:
#                 reader = pypdf.PdfReader(f)
#                 return ' '.join(p.extract_text() for p in reader.pages if p.extract_text())
#         elif ext == '.docx':
#             doc = docx.Document(file_path)
#             return ' '.join(p.text for p in doc.paragraphs)
#         elif ext in ['.xls', '.xlsx']:
#             df = pd.read_excel(file_path)
#             return df.to_string()
#         else:
#             raise ValueError("Unsupported file type.")
#     except Exception as e:
#         raise RuntimeError(f"Failed to process {file_path}: {str(e)}")


# # Utility: Ask Groq LLM a question
# def query_groq(context, question):
#     prompt = f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"
#     completion = groq_client.chat.completions.create(
#         messages=[{"role": "user", "content": prompt}],
#         model="llama3-70b-8192",
#     )
#     return completion.choices[0].message.content.strip()

# # Home page
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Upload files
# @app.route('/upload', methods=['POST'])
# def upload_file():
#     uploaded_files = request.files.getlist("file")
#     combined_text = ""

#     if not uploaded_files:
#         return jsonify({'error': 'No file uploaded'}), 400

#     filenames = []

#     for file in uploaded_files:
#         filename = secure_filename(file.filename)
#         ext = os.path.splitext(filename)[1].lower()

#         if ext not in ['.pdf', '.docx', '.xlsx', '.xls']:
#             return jsonify({'error': f'Unsupported file type: {ext}'}), 400

#         save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(save_path)
#         filenames.append(filename)

#         # Extract text
#         try:
#             file_text = process_file(save_path)
#             combined_text += "\n\n" + file_text
#         except Exception as e:
#             return jsonify({'error': str(e)}), 500

#     # Store in session for later queries
#     session['filenames'] = filenames
#     session['combined_text'] = combined_text

#     # Indexing
#     documents = [Document(text=combined_text)]
#     index = VectorStoreIndex.from_documents(documents, embed_model=embed_model)
#     temp_index_path = os.path.join(tempfile.gettempdir(), "index_tmp")
#     index.storage_context.persist(persist_dir=temp_index_path)
#     session['index_path'] = temp_index_path

#     return jsonify({'message': 'Files uploaded and processed', 'index_path': temp_index_path})

# # Query route
# @app.route('/query', methods=['POST'])
# def query_document():
#     data = request.get_json()
#     question = data.get('question', '')

#     if not question:
#         return jsonify({'error': 'Question is required'}), 400

#     combined_text = session.get('combined_text')
#     if not combined_text:
#         return jsonify({'error': 'No document uploaded yet'}), 400

#     try:
#         answer = query_groq(combined_text, question)
#         return jsonify({'response': answer})
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# # Summarize route
# @app.route('/summary', methods=['GET'])
# def summarize_document():
#     combined_text = session.get('combined_text')
#     if not combined_text:
#         return jsonify({'error': 'No document uploaded yet'}), 400

#     try:
#         summary_prompt = f"Please summarize the following document in a few concise bullet points:\n\n{combined_text}"
#         summary = query_groq(combined_text, summary_prompt)
#         return jsonify({'summary': summary})
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# # Delete uploaded files and session
# @app.route('/delete', methods=['POST'])
# def delete_files():
#     filenames = session.get('filenames', [])
#     for filename in filenames:
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         if os.path.exists(file_path):
#             os.remove(file_path)

#     session.clear()
#     return jsonify({'message': 'All uploaded files deleted successfully.'})

# # Run the app
# if __name__ == '__main__':
#     app.run(debug=True)
# app.py
import os
import tempfile
import logging
from pathlib import Path

from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, session
from werkzeug.utils import secure_filename

import pandas as pd
import docx
import pypdf

# LLM / vector imports (lazy load embeddings)
from llama_index.core import VectorStoreIndex, Document, StorageContext, load_index_from_storage, Settings
from llama_index.llms.gemini import Gemini

# -------------------------
# Config & Logging
# -------------------------
load_dotenv(".env.local")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("Missing GEMINI_API_KEY in .env.local")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# -------------------------
# Flask App Setup
# -------------------------
app = Flask(__name__)
app.secret_key = os.urandom(24)

app.config['UPLOAD_FOLDER'] = "uploads"
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500 MB
Path(app.config['UPLOAD_FOLDER']).mkdir(parents=True, exist_ok=True)

_embed_model = None  # will load when needed

def get_llm():
    return Gemini(model="models/gemini-1.5-flash", api_key=GEMINI_API_KEY)



def get_embed_model():
    """Lazy load embedding model."""
    global _embed_model
    if _embed_model is None:
        logger.info("Loading Gemini embedding model...")
        from llama_index.embeddings.gemini import GeminiEmbedding
        _embed_model = GeminiEmbedding(
            model_name="models/text-embedding-004",
            api_key=GEMINI_API_KEY
        )
        logger.info("Gemini Embedding model loaded.")
    return _embed_model


# -------------------------
# File Processing
# -------------------------
def process_file(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                return "\n".join([p.extract_text() or "" for p in reader.pages])
        elif ext == ".docx":
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        elif ext in [".xls", ".xlsx"]:
            df = pd.read_excel(file_path)
            return df.to_string(index=False)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        raise RuntimeError(f"Error processing {file_path}: {e}")


# -------------------------
# Routes
# -------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload_file():
    files = request.files.getlist("file")
    if not files:
        return jsonify({"error": "No file uploaded"}), 400

    combined_text_pieces = []
    filenames = []

    for file in files:
        filename = secure_filename(file.filename)
        if not filename:
            continue
        ext = os.path.splitext(filename)[1].lower()
        if ext not in [".pdf", ".docx", ".xlsx", ".xls"]:
            return jsonify({"error": f"Unsupported file: {ext}"}), 400

        save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(save_path)
        filenames.append(filename)

        try:
            text = process_file(save_path)
            if text:
                combined_text_pieces.append(text)
        except Exception as e:
            logger.exception("Error processing file: %s", filename)
            return jsonify({"error": str(e)}), 500

    if not combined_text_pieces:
        return jsonify({"error": "No text extracted"}), 400

    combined_text = "\n\n".join(combined_text_pieces)

    temp_text_fd, temp_text_path = tempfile.mkstemp(
        prefix="combined_", suffix=".txt"
    )
    os.close(temp_text_fd)
    with open(temp_text_path, "w", encoding="utf-8") as f:
        f.write(combined_text)

    session["filenames"] = filenames
    session["text_path"] = temp_text_path

    try:
        documents = [Document(text=combined_text)]
        index = VectorStoreIndex.from_documents(documents, embed_model=get_embed_model())

        temp_index_dir = tempfile.mkdtemp(prefix="llama_index_")
        index.storage_context.persist(persist_dir=temp_index_dir)
        session["index_path"] = temp_index_dir
    except Exception as e:
        logger.exception("Failed to build index")
        return jsonify({"error": f"Failed to create index: {e}"}), 500

    return jsonify({"message": "Files uploaded & processed"})


@app.route("/query", methods=["POST"])
def query_document():
    data = request.get_json(force=True)
    question = data.get("question", "").strip()
    if not question:
        return jsonify({"error": "Question is required"}), 400

    index_path = session.get("index_path")
    if not index_path or not os.path.exists(index_path):
        return jsonify({"error": "No document uploaded"}), 400

    try:
        storage_context = StorageContext.from_defaults(persist_dir=index_path)
        index = load_index_from_storage(storage_context, embed_model=get_embed_model())
        
        query_engine = index.as_query_engine(
            llm=get_llm(),
            similarity_top_k=5,
            response_mode="compact"
        )
        response = query_engine.query(question)
        
        # --- PROOF OF RAG ---
        # Log the exact chunks retrieved from the Vector Store before sending to Gemini
        logger.info(f"\n--- RAG RETRIEVAL: Found {len(response.source_nodes)} relevant chunks ---")
        for i, node in enumerate(response.source_nodes):
            logger.info(f"Chunk {i+1} (Score: {node.score:.4f}): {node.text[:150]}...")
        logger.info("--------------------------------------------------\n")
        
        return jsonify({"response": str(response)})
    except Exception as e:
        logger.exception("LlamaIndex query failed")
        return jsonify({"error": str(e)}), 500


@app.route("/summary", methods=["GET"])
def summarize_document():
    index_path = session.get("index_path")
    if not index_path or not os.path.exists(index_path):
        return jsonify({"error": "No document uploaded"}), 400

    summary_question = (
        "Please provide a comprehensive summary of this document. "
        "Include: 1) Main topics and themes, 2) Key findings or conclusions, "
        "3) Important data or statistics mentioned, 4) Any action items or recommendations. "
        "Format with clear markdown headings and bullet points."
    )

    try:
        storage_context = StorageContext.from_defaults(persist_dir=index_path)
        index = load_index_from_storage(storage_context, embed_model=get_embed_model())
        
        query_engine = index.as_query_engine(
            llm=get_llm(),
            response_mode="tree_summarize"
        )
        response = query_engine.query(summary_question)
        return jsonify({"summary": str(response)})
    except Exception as e:
        logger.exception("LlamaIndex summary failed")
        return jsonify({"error": str(e)}), 500


@app.route("/health", methods=["GET"])
def health_check():
    has_doc = bool(session.get("text_path") and os.path.exists(session.get("text_path", "")))
    filenames = session.get("filenames", [])
    return jsonify({"status": "ok", "has_document": has_doc, "filenames": filenames})


@app.route("/delete", methods=["POST"])
def delete_files():
    for filename in session.get("filenames", []):
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        if os.path.exists(file_path):
            os.remove(file_path)

    for path_key in ["text_path", "index_path"]:
        path = session.get(path_key)
        if path and os.path.exists(path):
            if os.path.isdir(path):
                for root, dirs, files in os.walk(path, topdown=False):
                    for name in files:
                        os.remove(os.path.join(root, name))
                    for name in dirs:
                        os.rmdir(os.path.join(root, name))
                os.rmdir(path)
            else:
                os.remove(path)

    session.clear()
    return jsonify({"message": "All uploaded files deleted"})


# -------------------------
# Entrypoint
# -------------------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
